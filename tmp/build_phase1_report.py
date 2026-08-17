from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = Path(r"C:\Users\DELL\OneDrive - Kathmandu BernHardt College\Desktop\Cards\docs\Tap2Connect_Phase_1_SDLC_Planning_Report.docx")
FIGJAM_URL = "https://www.figma.com/board/rtji9gPn2g3cT3F0qCHzv5"

BLUE = "2E74B5"
NAVY = "17365D"
INK = "111827"
MUTED = "52657A"
PALE_BLUE = "EAF2F8"
PALE_GREY = "F2F4F7"
TEAL = "117A65"
AMBER = "B9770E"
WHITE = "FFFFFF"


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margin(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_col_width(cell, width_inches):
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, text, url, color=BLUE):
    hyperlink = OxmlElement("w:hyperlink")
    part = paragraph.part
    relation_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_color = OxmlElement("w:color")
    r_color.set(qn("w:val"), color)
    r_pr.append(r_color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def text_run(paragraph, text, bold=False, color=None, size=None, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return run


def add_para(doc, text="", style=None, before=0, after=6, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if text:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = add_para(doc, style="List Bullet", after=3)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = add_para(doc, style="List Number", after=3)
        p.add_run(item)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_callout(doc, title, body, color=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    shade(cell, color)
    set_cell_margin(cell, 125, 170, 125, 170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    text_run(p, title, bold=True, color=NAVY)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    text_run(p, body, color=INK)
    add_para(doc, after=7)


def add_table(doc, headers, rows, widths=None, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, name in enumerate(headers):
        cell = header.cells[index]
        if widths:
            set_col_width(cell, widths[index])
        shade(cell, PALE_GREY)
        set_cell_margin(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        text_run(p, name, bold=True, color=NAVY, size=font_size)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            if widths:
                set_col_width(cell, widths[index])
            if row_index % 2 == 1:
                shade(cell, "FAFBFC")
            set_cell_margin(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            text_run(p, str(value), color=INK, size=font_size)
    add_para(doc, after=6)
    return table


def add_caption(doc, caption):
    p = add_para(doc, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    text_run(p, caption, italic=True, color=MUTED, size=9)


def page_break(doc):
    doc.add_page_break()


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color in (("Title", 30, NAVY), ("Subtitle", 15, MUTED), ("Heading 1", 16, BLUE), ("Heading 2", 13, NAVY), ("Heading 3", 11.5, NAVY)):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True if name != "Subtitle" else False
    styles["Heading 1"].paragraph_format.space_before = Pt(12)
    styles["Heading 1"].paragraph_format.space_after = Pt(6)
    styles["Heading 2"].paragraph_format.space_before = Pt(10)
    styles["Heading 2"].paragraph_format.space_after = Pt(4)
    for style_name in ("List Bullet", "List Number"):
        styles[style_name].font.name = "Calibri"
        styles[style_name].font.size = Pt(10.5)
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    text_run(p, "Tap2Connect Nepal  |  Phase 1 SDLC Planning", color=MUTED, size=8.5)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    text_run(p, "Confidential product planning draft  |  Version 1.0  |  August 2026", color=MUTED, size=8.5)


def build():
    doc = Document()
    configure_document(doc)

    # Cover page
    add_para(doc, before=36, after=8)
    p = add_para(doc, after=12)
    text_run(p, "TAP2CONNECT NEPAL", bold=True, color=BLUE, size=13)
    p = add_para(doc, after=8)
    text_run(p, "Smart Digital Identity Platform", bold=True, color=NAVY, size=29)
    p = add_para(doc, after=16)
    text_run(p, "Phase 1 SDLC Planning Report and System Blueprint", color=MUTED, size=16)
    add_callout(
        doc,
        "Planning purpose",
        "Define the product scope, users, data flows, architecture, delivery roadmap, controls, and success criteria for turning conventional identity cards into secure NFC and QR enabled digital identities.",
    )
    add_para(doc, after=5)
    meta = add_table(
        doc,
        ["Document control", "Details"],
        [
            ["Prepared for", "Tap2Connect Nepal product and implementation team"],
            ["Document type", "Phase 1: Planning and requirements baseline"],
            ["Version", "1.0"],
            ["Date", "14 August 2026"],
            ["Primary launch", "Organizations, schools, professionals, and individual users"],
        ],
        widths=[1.7, 4.9],
        font_size=10,
    )
    add_para(doc, after=50)
    p = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    text_run(p, "One physical card. A living digital identity.", bold=True, color=NAVY, size=14)
    page_break(doc)

    # Executive summary
    add_heading(doc, "1. Executive Summary", 1)
    add_para(doc, "Tap2Connect Nepal is a smart digital identity platform that converts ordinary printed identity cards, visiting cards, student cards, and professional cards into secure digital experiences. A tap on an NFC card or a scan of a QR code opens a mobile-first profile that is current, branded, and controlled by the profile owner or their organization.")
    add_para(doc, "The product serves two distinct but connected needs. The first is an organization contact-card experience for employees, staff, professionals, and business representatives. It prioritizes fast contact sharing, organization credibility, and a clean professional presence. The second is a student academic-resume experience. It begins with institutional identity but grows into a verified personal profile containing academic background, skills, projects, achievements, documents, and portfolio links.")
    add_callout(doc, "Core product decision", "Build one secure identity platform with two intentional template families—not two disconnected products. The account, role, analytics, QR/NFC mapping, permissions, and card-design engine are shared. The fields, display hierarchy, and management rules differ by template.", "E8F4F8")
    add_heading(doc, "1.1 Product Vision", 2)
    add_para(doc, "Make every identity card actionable. A card should no longer be a static piece of plastic or paper; it should be a controlled entry point to the right information, for the right audience, at the right time.")
    add_heading(doc, "1.2 Product Mission", 2)
    add_para(doc, "Enable organizations, schools, professionals, and individuals in Nepal to issue, manage, update, and share trustworthy digital identities through NFC and QR enabled cards without requiring recipients to install an application.")
    add_heading(doc, "1.3 Business Outcomes", 2)
    add_bullets(doc, [
        "Replace outdated, one-time printed information with a profile that can be updated after a card has been issued.",
        "Give organizations a controlled, brand-consistent contact-card solution for employees, departments, and representatives.",
        "Give students a verified, portable academic-resume profile that supports internships, admissions, networking, and employability.",
        "Give individual users a lightweight professional identity that is easier to share than a printed visiting card.",
        "Create a recurring platform relationship through card issuance, workspace subscriptions, design services, and managed identity operations.",
    ])

    add_heading(doc, "2. Problem Statement and Opportunity", 1)
    add_heading(doc, "2.1 Current Pain Points", 2)
    add_table(doc, ["Audience", "Current problem", "Tap2Connect response"], [
        ["Organizations", "Printed employee and visitor cards become outdated, do not show useful contact actions, and are hard to govern at scale.", "An organization workspace issues branded profiles, controls visible fields, assigns cards, and updates profiles centrally."],
        ["Schools and colleges", "Student IDs prove identity but rarely communicate skills, achievements, academic history, or portfolio work.", "A student academic-resume template connects verified identity with academic and employability information."],
        ["Professionals", "Paper business cards are lost; recipients must manually type contact information and cannot see work samples.", "A tap or scan opens a concise professional profile with one-tap contact, save-contact, portfolio, and social actions."],
        ["Recipients", "They receive fragmented information and must search separately for a phone number, email, website, or LinkedIn page.", "A public profile gives a fast, mobile-first view and safe actions without requiring sign-in."],
    ], widths=[1.25, 2.8, 2.65])
    add_heading(doc, "2.2 Product Opportunity", 2)
    add_para(doc, "Tap2Connect occupies the space between physical identity products and digital profile software. Its differentiator is not only NFC technology: it is the combination of physical card issuance, identity governance, role-based administration, profile templates, privacy controls, and measurable connection outcomes.")

    page_break(doc)
    add_heading(doc, "3. Scope and Product Boundaries", 1)
    add_heading(doc, "3.1 In-Scope Product Pillars", 2)
    add_table(doc, ["Pillar", "What it includes", "Why it matters"], [
        ["Identity workspaces", "Organization and personal workspaces, membership, roles, brand settings, and account recovery.", "Separates each organization’s data and lets administrators control who can issue or edit profiles."],
        ["Profile templates", "Organization contact template and student academic-resume template, with field groups, visibility rules, and reusable design variants.", "Makes the product fit different use cases without building a separate application for each."],
        ["Card issuance", "CR80 portrait and landscape designs, QR assignment, NFC token assignment, activation, replacement, suspension, and reissue status.", "Connects the physical card to a reliable digital identity."],
        ["Public sharing", "Public profile page, click-to-call, email, map, website, social links, vCard download, documents, and safe connection actions.", "Creates immediate value for the recipient after a tap or scan."],
        ["Administration and insight", "Bulk onboarding, approvals, profile status, card inventory, tap/scan activity, exports, reports, and audit history.", "Lets an organization operate the system instead of treating cards as one-off designs."],
    ], widths=[1.35, 3.45, 1.9])
    add_heading(doc, "3.2 Explicitly Out of Scope for the First Release", 2)
    add_bullets(doc, [
        "E-commerce shop, marketplace, or business-suite functions unrelated to smart identity cards.",
        "Payment processing and card purchase checkout; these can be evaluated after the platform workflow is stable.",
        "Native iOS or Android applications; the public profile and management portal are responsive web applications first.",
        "Unrestricted public editing of organization-owned identities.",
        "High-risk personal data collection not needed for the profile purpose, such as government document scans by default.",
        "AI-generated profile content as a mandatory dependency; AI may be introduced later as optional writing assistance.",
    ])
    add_heading(doc, "3.3 Product Principles", 2)
    add_numbered(doc, [
        "Fast first interaction: the recipient should see a usable profile in seconds after a tap or scan.",
        "Institutional control with individual dignity: organizations manage brand and access; owners manage their approved personal content.",
        "Privacy by design: public information is deliberate, minimum necessary, and reversible.",
        "One source of truth: profile changes update the digital card immediately; physical cards never need reprinting for ordinary profile updates.",
        "Mobile-first design: public pages, editing controls, and QR actions must work on a phone before being optimized for desktop.",
    ])

    add_heading(doc, "4. Template Strategy", 1)
    add_para(doc, "The templates are product experiences, not just visual themes. Both use the same platform infrastructure, but each prioritizes a different user goal and data model.")
    add_table(doc, ["Dimension", "Organization Contact Card", "Student Academic-Resume Card"], [
        ["Primary purpose", "Make a trusted professional connection and provide immediate contact details.", "Present a verified student identity plus a portable academic and employability profile."],
        ["Primary owner", "Organization, department, employer, or professional account.", "School/college plus the student, with configurable edit rights."],
        ["Audience", "Clients, partners, visitors, colleagues, recruiters, and event contacts.", "Recruiters, scholarship panels, institutions, mentors, internship providers, and peers."],
        ["Hero content", "Name, role, organization, photo, logo, short value statement, contact actions.", "Name, photo, student ID, program, institution, academic status, statement, and contact actions."],
        ["Extended content", "Department, office, service areas, location, website, LinkedIn, portfolio, appointment link.", "Academic history, GPA if enabled, skills, projects, certifications, achievements, activities, documents, portfolio links."],
        ["Governance", "Organization controls template, required fields, visibility rules, card lifecycle, and approvers.", "Institution controls verified fields; student may update approved personal sections subject to review policy."],
        ["Success event", "Recipient saves contact, starts a conversation, opens service link, or requests a connection.", "Recipient views academics/portfolio, downloads a resume/vCard, opens a project, or contacts the student."],
    ], widths=[1.22, 3.0, 2.48])
    add_callout(doc, "Recommended future extension", "Use a third ‘Individual Professional’ profile only as a configuration of the organization contact-card experience at launch. It avoids duplicate engineering while still allowing freelancers and alumni to have personal cards.", "FEF5E7")

    page_break(doc)
    add_heading(doc, "5. Stakeholders, Users, and Roles", 1)
    add_heading(doc, "5.1 Stakeholder Groups", 2)
    add_bullets(doc, [
        "Tap2Connect platform team: owns the SaaS platform, global support, operational controls, and product roadmap.",
        "Organization leadership: purchases or approves the service, defines brand and policy, and needs usable reporting.",
        "School or college administration: issues student identities, verifies information, manages bulk data, and protects student data.",
        "Profile owners: employees, professionals, students, teachers, and staff who keep their profile current within policy.",
        "Recipients: people who tap or scan and want useful contact or portfolio information immediately.",
        "Card production/NFC operator: prepares physical cards, binds QR/NFC identifiers, and manages replacement or activation states.",
    ])
    add_heading(doc, "5.2 Role-Based Access Model", 2)
    add_table(doc, ["Role", "Primary responsibilities", "Must not be able to"], [
        ["Platform Super Admin", "Manage tenants, platform settings, support cases, feature flags, audit access, and global templates.", "Change an organization’s public content without a traceable support reason and audit record."],
        ["Organization Admin", "Manage organization profile, brand, members, card templates, card assignments, approvals, and organization reports.", "Access any other organization’s data or bypass platform security controls."],
        ["School Admin", "Manage school branding, student intake, verified academic data, student-card policies, bulk upload, and approval queues.", "Publish private student records or allow unrestricted edits to verified academic fields."],
        ["Card/Content Manager", "Create and edit profiles, assign cards, prepare print/export batches, and review content where delegated.", "Change billing, tenant security, or super-admin settings."],
        ["Profile Owner", "Edit allowed personal information, links, skills, portfolio details, and requested updates; view personal analytics.", "Edit another owner’s profile, organization brand, or restricted academic identity fields."],
        ["Recipient / Visitor", "View public profile; save contact; open approved links; submit a connection request where enabled.", "See non-public information, internal analytics, edit controls, or other members’ profiles."],
    ], widths=[1.3, 3.45, 1.95])
    add_heading(doc, "5.3 Permission Design", 2)
    add_para(doc, "Every permission should be scoped by tenant (organization), resource type, action, and ownership. For example, ‘edit profile’ is not enough: the system must distinguish edit-own-profile, edit-member-profile, edit-verified-fields, publish-profile, issue-card, and view-organization-analytics. This prevents the current common problem where a broad ‘account’ role becomes unclear or unsafe.")

    add_heading(doc, "6. Functional Modules", 1)
    add_table(doc, ["Module", "Core functions", "Phase"], [
        ["Authentication and access", "Email/password or invited login, password reset, session handling, role lookup, account status, audit events.", "MVP"],
        ["Workspace and organization", "Organization/school profile, logo, colors, membership, departments, policies, brand defaults.", "MVP"],
        ["Profile management", "Create, edit, review, publish, archive, visibility controls, ownership transfer, and profile status.", "MVP"],
        ["Template and design studio", "Template families, field schemas, CR80 portrait/landscape formats, reusable layouts, assets, versions.", "MVP"],
        ["Student academic portfolio", "Academic background, skills, projects, certificates, achievement, document and portfolio blocks.", "MVP"],
        ["Card inventory and issuance", "Card stock, NFC/QR identifier binding, activation, revoke, replacement, print/batch export.", "MVP"],
        ["Public profile and share actions", "Slug/short URL, mobile profile, vCard, contact actions, open links, share, privacy-safe page.", "MVP"],
        ["Connections", "Recipient connection request, consent, accept/decline, basic inbox and notification state.", "Release 2"],
        ["Analytics and reporting", "Views, scans, taps, saves, interaction breakdown, member and institution reports.", "MVP basic; expand Release 2"],
        ["Imports and operations", "CSV bulk upload, validation preview, error report, card batch handling, export.", "MVP"],
    ], widths=[1.48, 4.35, 0.87], font_size=8.8)

    page_break(doc)
    add_heading(doc, "7. Data Flow Diagrams and Key Workflows", 1)
    add_para(doc, "The editable diagrams are maintained in FigJam so the product team can refine them during workshops. They represent the proposed target system, not merely the current development server.")
    p = add_para(doc, after=8)
    text_run(p, "Editable FigJam diagrams: ", bold=True, color=NAVY)
    add_hyperlink(p, "Tap2Connect Context DFD and Card Lifecycle", FIGJAM_URL)
    add_heading(doc, "7.1 Context DFD: System Boundary", 2)
    add_table(doc, ["External entity", "Data sent to Tap2Connect", "Data received from Tap2Connect"], [
        ["Organization / School Admin", "Organization details, brand configuration, member records, approval decisions, card issue requests.", "Dashboards, member status, card status, reports, error lists, audit information."],
        ["Student / Employee / Individual", "Profile details, permitted updates, links, documents, portfolio data, privacy choices.", "Published profile, edit status, feedback, card state, personal insights."],
        ["Public Recipient", "Card token or QR route, connection request, interaction action.", "Public profile, contact save file, approved contact and portfolio actions."],
        ["NFC Card / QR Channel", "Unique card or QR token that identifies a share route.", "Resolved, active public profile destination; no private data stored on the physical token."],
        ["Email Provider", "Invitation, password reset, connection or approval notification request.", "Delivery status and provider error response."],
    ], widths=[1.35, 3.1, 2.25])
    add_heading(doc, "7.2 Level 1 DFD: Core Processes", 2)
    add_table(doc, ["Process", "Input", "Processing", "Primary stores"], [
        ["P1. Identity and access", "Login, invite acceptance, password reset, session request.", "Authenticate, authorize role, create session, record security event.", "Users, roles, sessions, audit log."],
        ["P2. Workspace and membership", "Organization settings, roster, department, brand settings.", "Create tenant, apply policy, add membership, manage status.", "Organizations, memberships, brand settings."],
        ["P3. Profile and template", "Profile fields, template selection, visibility, documents, review decision.", "Validate schema, save draft, approve/publish, version design.", "Profiles, field values, templates, assets."],
        ["P4. Card issuance", "Card inventory item, QR/NFC token, assigned profile, activation request.", "Bind token to profile, generate route, set active/revoked/replaced status.", "Cards, tokens, issue batches."],
        ["P5. Public share", "Tap/scan route, public action.", "Validate card state, resolve profile, apply visibility policy, render action.", "Cards, published profiles, action events."],
        ["P6. Insight and operations", "Events, export request, report filter.", "Aggregate metrics, produce tenant-safe reports and operational exports.", "Activity events, audit logs, reports."],
    ], widths=[1.3, 1.55, 2.75, 1.1], font_size=8.8)
    add_heading(doc, "7.3 Card Lifecycle Flow", 2)
    add_numbered(doc, [
        "A platform administrator or customer administrator creates a workspace and chooses organization, school, or individual onboarding.",
        "The user creates or imports a profile and chooses the appropriate template family.",
        "The system validates required fields, applies organization visibility rules, and routes the profile for review if policy requires it.",
        "An approved profile is assigned a card record and a QR/NFC token. The physical card stores only the safe identifier or URL token—not personal profile data.",
        "A recipient taps or scans; the platform checks whether the card is active, resolves the route, and displays the public profile.",
        "The recipient saves contact details, opens an approved link, downloads a vCard, or requests a connection. Each event is logged with a privacy-safe activity record.",
        "The owner or administrator updates the profile. The next tap or scan shows the latest published version without reprinting the card.",
    ])

    add_heading(doc, "8. Target System Architecture", 1)
    add_para(doc, "The target architecture should remain a modular monolith during MVP: a React web frontend and a Django REST backend with clear service boundaries. This is simpler to build, test, and operate than early microservices, while still allowing later separation of analytics, notifications, or card provisioning if scale requires it.")
    add_table(doc, ["Layer", "Recommended responsibility", "Implementation direction"], [
        ["Presentation", "Responsive public profiles, admin dashboard, profile editor, template studio, card designer.", "React + Vite today; preserve component library and token-based design system."],
        ["Application API", "Authentication, role checks, tenant isolation, validation, workflows, reporting endpoints.", "Django + Django REST Framework style APIs; use a consistent /api/v1 contract before public integrations."],
        ["Domain services", "Profiles, organization membership, templates, cards, share resolution, analytics, imports.", "Separate Django app/service modules with explicit ownership and tests."],
        ["Data", "Relational records, media assets, audit events, template versions, interaction logs.", "PostgreSQL for production. SQLite is acceptable only for local development."],
        ["Asset storage", "Profile photos, logos, PDFs, certificate previews, card-design assets.", "Private object storage with signed access for private media; public CDN only for public assets."],
        ["External channels", "NFC card URL/token, QR generation, email delivery, optional SMS/WhatsApp notification later.", "Token-based routes, provider adapters, delivery logs, and no private profile data on the card."],
    ], widths=[1.25, 3.3, 2.15])
    add_callout(doc, "Architecture decision", "Treat ‘College’ as a type of Organization in the target data model. The UI can still say School or College, but the database should avoid duplicating the whole organization concept for each vertical.", "FEF5E7")

    page_break(doc)
    add_heading(doc, "9. Target Data Model", 1)
    add_para(doc, "The data model should be designed around tenancy, identity ownership, reusable templates, and card lifecycle. The exact field list may evolve, but the relationship boundaries should be agreed before the production database is finalized.")
    add_table(doc, ["Entity", "Purpose", "Key relationships"], [
        ["User", "Login identity for a platform administrator, organization user, or profile owner.", "Has memberships, owns permitted profile edits, produces audit events."],
        ["Organization", "Tenant for a company, school, college, NGO, or partner organization.", "Has memberships, brand settings, templates, profiles, cards, policies."],
        ["Membership", "Role assignment connecting a user to an organization.", "Belongs to user and organization; carries role and status."],
        ["Identity Profile", "The live digital identity that a recipient sees.", "Belongs to organization optionally; has template, visibility, owner, status, and card assignments."],
        ["Profile Field / Section", "Structured data such as contact, academic background, skills, achievements, or project details.", "Belongs to profile and template schema; supports repeatable sections."],
        ["Card Template", "A reusable visual and field schema for a profile family.", "Has versions and allowed field definitions; used by many profiles/cards."],
        ["Card", "Physical or virtual issued item that resolves to one identity profile.", "Belongs to organization/card batch; references active profile and token; has lifecycle state."],
        ["Share Token", "Opaque QR/NFC route token with no private payload.", "Belongs to card; maps to active public route and supports revoke/replace."],
        ["Asset / Document", "Image, logo, certificate, resume, project file, or design asset.", "Owned by organization/profile; visibility and access policy are explicit."],
        ["Activity Event", "Privacy-safe tap, scan, view, download, action, or connection event.", "References profile/card/organization; used for analytics with retention policy."],
        ["Connection Request", "Consent-based request made from a public profile.", "Has recipient contact data, profile target, state, and notification history."],
    ], widths=[1.4, 2.8, 2.5], font_size=8.8)
    add_heading(doc, "9.1 Required Card States", 2)
    add_table(doc, ["State", "Meaning", "Public behaviour"], [
        ["Draft", "Profile/card is being prepared and has not been approved or bound.", "Not resolvable."],
        ["Pending approval", "Profile content or card issuance is waiting for authorized review.", "Not resolvable unless a private preview token is used."],
        ["Active", "Card is issued and public profile is published.", "Resolves to public profile."],
        ["Suspended", "Temporarily disabled due to policy, loss, departure, or review.", "Shows a neutral unavailable page; does not expose old details."],
        ["Replaced", "A new card supersedes this physical token.", "Shows a neutral replacement/unavailable result or redirects only if policy permits."],
        ["Archived", "Historical record retained for audit but no longer in use.", "Not resolvable."],
    ], widths=[1.45, 2.9, 2.35])
    add_heading(doc, "9.2 Field Catalog", 2)
    add_table(doc, ["Field group", "Organization contact template", "Student academic-resume template"], [
        ["Identity", "Name, photo, job title, department, organization, employee/reference ID.", "Name, photo, student ID, institution, program, cohort, graduation year."],
        ["Contact", "Phone, office email, website, office address, map, appointment link.", "Phone/email if enabled, guardian/emergency contact only if strictly required and never public by default."],
        ["Professional / academic", "Service areas, biography, expertise, LinkedIn, portfolio.", "Education, course/program, academic highlights, skills, projects, certifications, languages, achievements."],
        ["Actions", "Save contact, call, email, map, website, WhatsApp, social, portfolio.", "Save contact, resume/portfolio, project links, email, social, verified document links where allowed."],
        ["Privacy", "Admin-controlled contact/action visibility.", "Institution-controlled verified data, student-owned personal fields, per-section public/private settings."],
    ], widths=[1.25, 3.15, 2.3], font_size=8.8)

    page_break(doc)
    add_heading(doc, "10. Functional Requirements", 1)
    add_heading(doc, "10.1 MVP Requirements", 2)
    requirements = [
        ("FR-01", "The system shall support organization, school/college, and individual workspace onboarding."),
        ("FR-02", "The system shall support role-based access with tenant-scoped permissions and profile ownership rules."),
        ("FR-03", "An administrator shall be able to create, import, edit, approve, suspend, archive, and assign identity profiles."),
        ("FR-04", "The system shall provide two reusable template families: organization contact and student academic-resume."),
        ("FR-05", "A template shall define required fields, allowed field groups, display order, visibility rules, orientation, and design version."),
        ("FR-06", "The system shall support CR80 portrait (54 × 86 mm) and landscape (86 × 54 mm) card formats."),
        ("FR-07", "The system shall assign each issued card a unique QR and/or NFC share token that can be activated, suspended, replaced, or revoked."),
        ("FR-08", "The public profile shall work without login and offer only fields/actions approved for public visibility."),
        ("FR-09", "A recipient shall be able to save a contact as a standards-compliant vCard and use permitted call, email, map, web, and social actions."),
        ("FR-10", "The system shall capture profile views and interaction events and show tenant-safe analytics to authorized users."),
        ("FR-11", "School/organization administrators shall be able to import members in bulk with validation, preview, and error reporting."),
        ("FR-12", "The system shall maintain audit records for publish, visibility, role, card-state, and sensitive field changes."),
    ]
    add_table(doc, ["ID", "Requirement"], requirements, widths=[0.75, 6.0], font_size=9)
    add_heading(doc, "10.2 Release 2 Requirements", 2)
    add_bullets(doc, [
        "Consent-based connection workflow with notifications, accept/decline, and contact sharing preferences.",
        "Advanced template editor with searchable icon library, editable icon style, alignment tools, and template locking by organization policy.",
        "Card inventory batches and activation station workflow for physical production teams.",
        "Approval workflows by department, school faculty, or program coordinator.",
        "Expanded analytics: campaign/event attribution, activity trends, and organization-level CSV exports.",
        "Optional multilingual field labels and public profile localization (English and Nepali).",
    ])
    add_heading(doc, "10.3 Non-Functional Requirements", 2)
    add_table(doc, ["Area", "Requirement baseline"], [
        ["Security", "HTTPS in production, secure cookies or token strategy, CSRF protection where session auth is used, strong password policy, rate-limited auth, and server-side authorization on every protected endpoint."],
        ["Privacy", "Public visibility is explicit by field/section. NFC/QR tokens carry no private profile data. Data retention and deletion rules are documented per tenant."],
        ["Performance", "Public card route target: usable first content within 2.5 seconds on a typical mobile network; cache safe public assets."],
        ["Availability", "Production backups, restore test, monitoring, error tracking, and an unavailable-card fallback page."],
        ["Accessibility", "Keyboard support, readable contrast, visible labels, semantic action buttons, alt text, and responsive layout."],
        ["Auditability", "Role changes, publication, field visibility, export, card activation, suspension, and replacement are traceable."],
        ["Data isolation", "Every organization query must be tenant filtered; no object may be loaded by ID alone without authorization context."],
    ], widths=[1.35, 5.4], font_size=9)

    page_break(doc)
    add_heading(doc, "11. Security, Privacy, and Compliance Plan", 1)
    add_heading(doc, "11.1 Security Controls", 2)
    add_bullets(doc, [
        "Use a production database (PostgreSQL) and protected environment variables; never commit secrets or production credentials to the repository.",
        "Use a single cross-origin authentication strategy. If React and Django use different origins during development, configure CORS and CSRF deliberately, and test login/connection endpoints end-to-end.",
        "Enforce server-side object permissions; frontend controls improve usability but do not provide security.",
        "Store password hashes only; protect admin accounts with stronger controls and plan for two-factor authentication before broad organization rollout.",
        "Use opaque, non-sequential public slugs/tokens for card routes where possible; revoke a card token immediately when a card is lost or a member leaves.",
        "Validate uploads by size/type, scan files in production, store private documents outside public media paths, and issue time-limited access URLs.",
        "Protect against common web risks: CSRF, XSS, injection, unsafe redirect URLs, brute-force logins, and insecure direct object references.",
    ])
    add_heading(doc, "11.2 Student Data Rules", 2)
    add_para(doc, "Student profiles need stricter defaults than ordinary contact cards. Academic results, guardian details, birth records, personal addresses, and private documents must never become public automatically. The institution should choose the verified fields it owns; the student should control optional personal portfolio fields; public display should be section-based and visible in a preview before publication.")
    add_heading(doc, "11.3 Data Classification", 2)
    add_table(doc, ["Classification", "Examples", "Handling"], [
        ["Public", "Name, role, approved photo, approved professional contact, public portfolio links.", "Visible on profile by published policy."],
        ["Organization internal", "Employee identifier, department management notes, card batch data, non-public reports.", "Visible only to appropriate tenant roles."],
        ["Sensitive", "Student private address, guardian contact, academic records, private certificates, login data.", "Private by default; least-privilege access, audit events, retention policy."],
        ["Security confidential", "Passwords, reset tokens, secrets, session data, API credentials.", "Never exposed in UI/export/logs; encrypted/hashed per purpose."],
    ], widths=[1.3, 3.1, 2.35])

    add_heading(doc, "12. Integration and API Plan", 1)
    add_table(doc, ["Integration", "MVP decision", "Notes"], [
        ["NFC cards", "Use a short HTTPS URL or opaque token written to compatible NFC tags/cards.", "The tag should contain only the route/token. Server state decides what happens after a tap."],
        ["QR codes", "Generate per-card QR routes with an error-correction level suitable for print.", "Support reissue/revoke without reprinting a full profile; route remains server-controlled."],
        ["vCard", "Generate contact card download from the currently published public fields.", "Avoid inserting non-public data into downloadable contact files."],
        ["Email", "Use a transactional provider for invites, reset passwords, approvals, and connection notifications.", "Store delivery status; provider adapter keeps future switching possible."],
        ["Maps/social/portfolio", "Permit approved external links with URL validation and optional organization allow-list.", "Links are public actions; do not proxy untrusted content."],
        ["Printing", "Export card-ready QR/token payloads and template proofs for the production partner.", "Physical printing is an operational integration, not a separate commerce module."],
    ], widths=[1.3, 2.8, 2.65], font_size=8.8)

    page_break(doc)
    add_heading(doc, "13. SDLC Delivery Plan", 1)
    add_para(doc, "This plan assumes a small product team (product/UX lead, React developer, Django developer, QA support, and card-production stakeholder) working in short, reviewable increments. The durations are planning estimates; the pilot should begin only after the defined acceptance criteria are met.")
    add_table(doc, ["Phase", "Duration", "Key outputs", "Exit criteria"], [
        ["0. Discovery and validation", "2 weeks", "Stakeholder interviews, field catalog, privacy policy, pilot organization, success metrics, clickable UX flows.", "Product brief, roles, template fields, and pilot commitments approved."],
        ["1. Foundation", "2 weeks", "Tenant model, authentication approach, React-Django API contract, PostgreSQL plan, design tokens, test environment.", "Login, tenant scoping, and baseline deployment work end-to-end."],
        ["2. Organization contact MVP", "3 weeks", "Organization workspace, member profile, contact template, public profile, QR, vCard, basic analytics.", "Admin can issue a contact card and a recipient can tap/scan/save safely."],
        ["3. Student academic-resume MVP", "3 weeks", "Student schema, institution verification, academic/skills/projects sections, privacy controls, approval workflow.", "School can import/approve a student profile; only approved fields are public."],
        ["4. Card operations", "2 weeks", "Card token bind/revoke/replace, batch export, CR80 orientation, print-production checks.", "A lost/replaced card is controlled instantly and a production batch is traceable."],
        ["5. Quality and pilot", "2 weeks", "Automated tests, accessibility pass, security review, monitoring, backups, UAT, pilot training.", "All MVP acceptance tests pass; pilot users complete real workflows."],
        ["6. Pilot and hardening", "2–4 weeks", "One organization/school pilot, feedback, support fixes, performance improvements, launch decision.", "Agreed activation, scan, completion, and issue thresholds are met."],
    ], widths=[1.45, 0.85, 3.0, 1.45], font_size=8.5)
    add_heading(doc, "13.1 Delivery Cadence", 2)
    add_bullets(doc, [
        "Weekly product review: validate templates, field choices, visibility rules, and pilot feedback with real users.",
        "Each sprint: define acceptance criteria before coding, demo working functionality, record defects, and update the risk register.",
        "Before changing a data model: review migration plan, tenant effects, API backward compatibility, and privacy impact.",
        "Before release: run backend unit/API tests, frontend build, responsive checks, login/CSRF/CORS checks, and public profile tests on a phone.",
    ])
    add_heading(doc, "13.2 MVP Backlog Priority", 2)
    add_table(doc, ["Priority", "Scope"], [
        ["Must", "Tenant isolation, roles, two template families, public profile, QR/NFC route resolution, vCard, profile visibility, card states, bulk onboarding, audit basics, responsive UI."],
        ["Should", "NFC activation station, connection requests, advanced editor, robust reports, notification delivery history, template version rollback."],
        ["Could", "AI writing assistant, multilingual public profiles, appointment calendars, CRM sync, digital wallet passes, custom domains."],
        ["Not now", "General e-commerce/business-suite features, unrelated online shop, complex marketplace, native apps before web product-market fit."],
    ], widths=[1.1, 5.7])

    page_break(doc)
    add_heading(doc, "14. Quality Assurance and Acceptance Criteria", 1)
    add_heading(doc, "14.1 Essential Acceptance Scenarios", 2)
    add_table(doc, ["Scenario", "Expected result"], [
        ["Organization admin creates a profile", "Profile is created in only that organization, required fields are validated, and audit entry is recorded."],
        ["Student edits profile", "Student can edit only permitted personal sections; verified academic fields follow school policy and approval rules."],
        ["Card is tapped", "Active token resolves to the correct public profile on mobile; suspended/replaced token does not expose prior details."],
        ["Recipient saves contact", "vCard contains only the currently published contact data and downloads successfully across common mobile browsers."],
        ["User logs in from React frontend", "Login, CSRF/CORS/session configuration works consistently from the deployed frontend origin; protected API calls succeed only when authorized."],
        ["Bulk upload runs", "Invalid records are reported clearly; valid records are staged safely; no records leak into another tenant."],
        ["Admin views analytics", "Metrics are scoped to the authorized organization and do not expose raw recipient personal information unnecessarily."],
        ["Lost card is suspended", "Future taps show a neutral unavailable state immediately; replacement has a new token and audit trail."],
    ], widths=[2.15, 4.65])
    add_heading(doc, "14.2 Test Layers", 2)
    add_bullets(doc, [
        "Unit tests for model validation, token resolution, permissions, vCard generation, and state transitions.",
        "API integration tests for login, CSRF/CORS configuration, organization filters, public profile visibility, and bulk imports.",
        "Frontend component tests for editor controls, responsive templates, error states, and role-based navigation.",
        "End-to-end browser tests for organization onboarding, student profile publishing, tap/scan route, and connection flow.",
        "Manual device testing across Android and iPhone NFC/QR behaviours, especially the public page and contact save experience.",
        "Security review before pilot, including tenant-bound object access and unauthenticated public route probing.",
    ])

    add_heading(doc, "15. Risks and Mitigations", 1)
    add_table(doc, ["Risk", "Impact", "Mitigation"], [
        ["Unclear template scope", "The editor becomes complex and each customer requests a custom product.", "Lock the two template families and field catalog for MVP; allow configuration only through approved schemas."],
        ["Mixed authentication origins", "403/CSRF/CORS failures interrupt login, connections, and editing.", "Document one production origin strategy; automate smoke tests from React to Django API before each release."],
        ["Weak tenant checks", "Organization data could be visible or editable across boundaries.", "Tenant filter every query, test authorization negatively, and centralize object-permission helpers."],
        ["Student privacy exposure", "Sensitive educational/personal information could be published accidentally.", "Private-by-default fields, preview, approval, audit, and restricted documents."],
        ["Lost or duplicated card", "Old card may continue to expose an identity.", "Token lifecycle with immediate suspend/replace/revoke and no sensitive payload in the physical card."],
        ["Overdesign before validation", "Time is spent on 3D visuals or advanced editors before the sharing workflow works.", "Prioritize card-to-profile value, mobile responsiveness, and pilot evidence; add visual polish after core flows are reliable."],
        ["Data import quality", "School/organization data is incomplete or inconsistent.", "Staging, validation report, templates, duplicate detection, and admin approval before publishing."],
    ], widths=[1.7, 2.25, 2.85], font_size=8.8)

    page_break(doc)
    add_heading(doc, "16. Success Metrics", 1)
    add_table(doc, ["Area", "Metric", "Pilot target / decision use"], [
        ["Activation", "Percentage of issued cards activated with a published profile.", "At least 90% of a pilot batch has a working public route before handover."],
        ["Profile completion", "Percentage of profile owners with required and recommended fields complete.", "At least 80% complete for pilot cohort; identify blocking fields."],
        ["Sharing", "Taps/scans per active card and public profile load success rate.", "Track adoption and infrastructure reliability; no unsupported assumptions about marketing conversion."],
        ["Action value", "vCard downloads, calls, emails, portfolio opens, and connection requests per public visit.", "Identify whether each template creates its intended outcome."],
        ["Operations", "Time to issue, update, suspend, and replace a card.", "Measure whether the organization can run the system without developer intervention."],
        ["Support quality", "Authentication issues, failed share routes, imports rejected, and median resolution time.", "Use to prioritize reliability work before scaling."],
        ["Privacy/security", "Unauthorized-access defects, accidental public-field findings, and audit coverage.", "Zero critical tenant/privacy defects at launch."],
    ], widths=[1.2, 3.5, 2.1])

    add_heading(doc, "17. Decisions Required Before Build Commitment", 1)
    add_numbered(doc, [
        "Choose the first pilot vertical: one organization contact-card rollout, one school/college rollout, or a controlled combination. Recommendation: run both with small cohorts because their templates validate different product value.",
        "Approve the MVP field catalog and decide which student fields are verified, student-editable, and public by default.",
        "Confirm the card production workflow: card supplier, NFC technology/tag type, QR print process, card inventory ownership, and replacement policy.",
        "Select the production hosting, database, media-storage, and transactional-email provider before real user data is loaded.",
        "Approve the brand system: logo, color tokens, type scale, imagery direction, and accessible public-page design before templates are locked.",
        "Assign data governance owners at each pilot organization: one business owner, one content/verification owner, and one support contact.",
    ])

    add_heading(doc, "18. Recommended Immediate Next Steps", 1)
    add_numbered(doc, [
        "Run a 60–90 minute requirements workshop using the attached DFD and lifecycle flow; finalize roles, field visibility, and card-state rules.",
        "Convert the two templates into wireframes and validate them with one organization representative, one school administrator, two students, and two card recipients.",
        "Refactor the target data model around Organization, Membership, IdentityProfile, Card, and ShareToken before adding more UI features.",
        "Stabilize React-to-Django authentication on a single documented dev/prod origin model; add tests for login, connection request, and protected profile edits.",
        "Build the organization contact-card MVP slice end-to-end first: create profile → publish → assign QR/NFC token → tap → save contact → report activity.",
        "Build the student academic-resume slice next, with privacy/approval rules before adding advanced visual editor capabilities.",
        "Pilot with a limited, named cohort; track the metrics in Section 16 and decide what deserves Release 2 investment.",
    ])

    add_heading(doc, "Appendix A. Diagram Legend", 1)
    add_table(doc, ["Term", "Meaning"], [
        ["Tenant", "A protected organization or school workspace whose users, profiles, cards, and reports are isolated from other customers."],
        ["Public profile", "The recipient-facing page shown after a valid NFC tap or QR scan. It exposes only published fields/actions."],
        ["Share token", "An opaque identifier held by the QR code or NFC card that is resolved by the server. It does not contain private data."],
        ["Template family", "A reusable product experience and field schema: organization contact or student academic-resume."],
        ["Card lifecycle", "The controlled sequence from draft to approval to active use, suspension, replacement, and archive."],
        ["Activity event", "A privacy-safe record of a view, scan, tap, contact download, approved link action, or connection action."],
    ], widths=[1.4, 5.4])
    add_para(doc, "End of report.", after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
